import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import io
from datetime import datetime

# ─── PAGE CONFIG ──────────────────────────────────────────────
st.set_page_config(
    page_title="GrantMaster Pro",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─── SESSION STATE INIT ────────────────────────────────────────
if "selected_grant" not in st.session_state:
    st.session_state.selected_grant = None
if "result_text" not in st.session_state:
    st.session_state.result_text = None
if "score_text" not in st.session_state:
    st.session_state.score_text = None
if "advice_text" not in st.session_state:
    st.session_state.advice_text = None
if "org_name" not in st.session_state:
    st.session_state.org_name = ""
if "prefill" not in st.session_state:
    st.session_state.prefill = {}

# ─── СТИЛЬ ────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display&family=DM+Sans:wght@300;400;500;600&family=JetBrains+Mono:wght@400;600&display=swap');

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
}

/* Root background */
.stApp {
    background: #0B0F1A;
    color: #E8EAF0;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background: #0D1422 !important;
    border-right: 1px solid #1E2A40;
}
section[data-testid="stSidebar"] * {
    color: #C5CDE0 !important;
}

/* Скрыть лишние элементы Streamlit */
#MainMenu, footer, header {visibility: hidden;}
.block-container {padding-top: 2rem; padding-bottom: 4rem; max-width: 1280px;}

/* Заголовки */
h1 { font-family: 'DM Serif Display', serif !important; color: #FFFFFF !important; font-size: 2.6rem !important; letter-spacing: -0.03em; }
h2 { font-family: 'DM Serif Display', serif !important; color: #C9A84C !important; font-size: 1.5rem !important; }
h3 { color: #8B9CC8 !important; font-size: 1rem !important; font-weight: 600 !important; letter-spacing: 0.06em; text-transform: uppercase; }

/* Карточки грантов */
.grant-card {
    background: #111826;
    border: 1px solid #1E2A40;
    border-radius: 12px;
    padding: 20px 18px;
    margin-bottom: 8px;
    transition: border-color 0.2s, box-shadow 0.2s;
    cursor: pointer;
}
.grant-card:hover {
    border-color: #C9A84C;
    box-shadow: 0 0 0 1px #C9A84C22, 0 8px 24px #00000044;
}
.grant-card.active {
    border-color: #C9A84C;
    background: #17202E;
    box-shadow: 0 0 0 2px #C9A84C44;
}
.grant-tag {
    display: inline-block;
    background: #1E2A40;
    color: #8B9CC8;
    font-size: 11px;
    font-weight: 600;
    padding: 3px 8px;
    border-radius: 4px;
    letter-spacing: 0.04em;
    margin-bottom: 6px;
}
.grant-name { font-size: 15px; font-weight: 600; color: #E8EAF0; margin-bottom: 4px; }
.grant-amount { font-size: 13px; color: #C9A84C; font-weight: 600; }
.grant-desc { font-size: 12px; color: #607090; margin-top: 4px; }

/* Форма */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stNumberInput > div > div > input {
    background: #111826 !important;
    border: 1px solid #1E2A40 !important;
    color: #E8EAF0 !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: #C9A84C !important;
    box-shadow: 0 0 0 1px #C9A84C44 !important;
}
label { color: #8B9CC8 !important; font-size: 12px !important; font-weight: 600 !important; letter-spacing: 0.05em !important; text-transform: uppercase !important; }

/* Кнопки */
.stButton > button {
    background: #C9A84C !important;
    color: #0B0F1A !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 700 !important;
    font-size: 13px !important;
    letter-spacing: 0.04em !important;
    padding: 10px 20px !important;
    transition: opacity 0.2s !important;
}
.stButton > button:hover { opacity: 0.85 !important; }
.stButton > button[kind="secondary"] {
    background: #1E2A40 !important;
    color: #8B9CC8 !important;
}

/* Download кнопки */
.stDownloadButton > button {
    background: #1E2A40 !important;
    color: #E8EAF0 !important;
    border: 1px solid #2A3A58 !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 13px !important;
    width: 100% !important;
}
.stDownloadButton > button:hover {
    border-color: #C9A84C !important;
    color: #C9A84C !important;
    background: #17202E !important;
}

/* Результат */
.result-box {
    background: #0D1422;
    border: 1px solid #1E2A40;
    border-radius: 12px;
    padding: 28px 32px;
    font-family: 'DM Sans', sans-serif;
    font-size: 14px;
    line-height: 1.8;
    color: #C5CDE0;
    white-space: pre-wrap;
}

/* Score */
.score-box {
    background: linear-gradient(135deg, #111826 0%, #0D1422 100%);
    border: 1px solid #C9A84C44;
    border-radius: 12px;
    padding: 24px;
    text-align: center;
}
.score-number {
    font-family: 'DM Serif Display', serif;
    font-size: 64px;
    color: #C9A84C;
    line-height: 1;
}
.score-label { font-size: 11px; color: #607090; letter-spacing: 0.08em; text-transform: uppercase; margin-top: 6px; }

/* Advice chips */
.advice-item {
    background: #111826;
    border: 1px solid #1E2A40;
    border-left: 3px solid #C9A84C;
    border-radius: 8px;
    padding: 14px 16px;
    margin-bottom: 10px;
    font-size: 13px;
    color: #C5CDE0;
    line-height: 1.6;
}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    background: #0D1422 !important;
    border-bottom: 1px solid #1E2A40 !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    color: #607090 !important;
    border-radius: 0 !important;
    padding: 12px 24px !important;
    font-weight: 600 !important;
    font-size: 13px !important;
    letter-spacing: 0.04em !important;
    border-bottom: 2px solid transparent !important;
}
.stTabs [aria-selected="true"] {
    color: #C9A84C !important;
    border-bottom: 2px solid #C9A84C !important;
    background: transparent !important;
}
.stTabs [data-baseweb="tab-panel"] {
    background: transparent !important;
    padding-top: 24px !important;
}

/* Divider */
hr { border-color: #1E2A40 !important; }

/* Alerts */
.stAlert { border-radius: 8px !important; border: none !important; }

/* Spinner */
.stSpinner > div { border-top-color: #C9A84C !important; }

/* Progress */
.stProgress > div > div > div { background: #C9A84C !important; }
</style>
""", unsafe_allow_html=True)

# ─── ДАННЫЕ ГРАНТОВ ────────────────────────────────────────────
GRANTS = {
    "Ścieżka SMART FENG": {
        "tag": "🇵🇱 POLSKA",
        "amount": "До 25 млн PLN",
        "desc": "Innowacje dla MŚP i dużych firm. R&D + wdrożenie.",
        "lang": "POLISH",
        "example": {
            "org": "InnoTech Sp. z o.o., Warszawa",
            "desc": "Platforma AI do optymalizacji łańcuchów dostaw w logistyce e-commerce. Redukcja kosztów o 35%, 10 nowych miejsc pracy w regionie mazowieckim.",
            "team": "CEO (MBA, 10 lat exp.) + 3 dev (ML/Python) + ekspert NCBR",
            "budget": "450000"
        }
    },
    "Horizon Europe EIC": {
        "tag": "🇪🇺 EU",
        "amount": "До 2.5 млн €",
        "desc": "Deep tech & breakthrough innovation for scale-ups.",
        "lang": "ENGLISH",
        "example": {
            "org": "DeepMed Technologies, Warsaw (EU-registered SME)",
            "desc": "Novel AI-driven diagnostics platform for early detection of rare diseases using federated learning. 60% accuracy improvement vs. current standard.",
            "team": "CTO (PhD, MIT) + 2 AI researchers + medical advisor",
            "budget": "1200000"
        }
    },
    "PARP dla МСП": {
        "tag": "🇵🇱 PARP",
        "amount": "До 1 млн PLN",
        "desc": "Internacjonalizacja i cyfryzacja małych firm.",
        "lang": "POLISH",
        "example": {
            "org": "GreenFlow Sp. z o.o., Kraków",
            "desc": "SaaS do zarządzania energią dla małych przedsiębiorstw. Oszczędności 20-40% kosztów energii, ESG compliance.",
            "team": "2 założycieli + developer + konsultant UE",
            "budget": "320000"
        }
    },
    "єРобота (Україна)": {
        "tag": "🇺🇦 UKRAINE",
        "amount": "До 250 000 ₴",
        "desc": "Підтримка малого бізнесу та самозайнятих.",
        "lang": "UKRAINIAN",
        "example": {
            "org": "ТОВ 'ТехноІнвест', Київ",
            "desc": "Платформа онлайн-навчання для ветеранів АТО та ООС. Охоплення 5000+ осіб, партнерство з МОН України.",
            "team": "Засновник + 2 методисти + технічний директор",
            "budget": "180000"
        }
    },
    "NCBR Strategiczne": {
        "tag": "🇵🇱 NCBR",
        "amount": "До 15 млн PLN",
        "desc": "Badania przemysłowe i prace rozwojowe — consortia.",
        "lang": "POLISH",
        "example": {
            "org": "Konsorcjum: PoliWarszawa + BioTech Sp. z o.o.",
            "desc": "Biodegradowalne opakowania z odpadów rolniczych (B2). Zastąpienie plastiku w 30% rynku FMCG. Ochrona IP — zgłoszony patent EP.",
            "team": "Lider naukowy (prof. dr hab.) + 5 badaczy + partner przemysłowy",
            "budget": "8500000"
        }
    },
    "МGO Ukraine (EU-funded)": {
        "tag": "🇺🇦🇪🇺 NGO",
        "amount": "До 500 000 €",
        "desc": "EU-funded grants for Ukrainian civil society & recovery.",
        "lang": "UKRAINIAN",
        "example": {
            "org": "ГО 'Центр Відновлення', Харків",
            "desc": "Програма психологічної реабілітації та працевлаштування для ВПО. 2000 бенефіціарів, партнерство з UNHCR та муніципалітетом.",
            "team": "Директор + 4 психологи + менеджер проектів ЄС",
            "budget": "280000"
        }
    },
}

# ─── SIDEBAR ──────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚡ GrantMaster Pro")
    st.caption("AI Grant Application Generator · 2026")
    st.divider()

    # API Key
    api_key_input = st.text_input(
        "Gemini API Key",
        type="password",
        placeholder="AIzaSy...",
        help="Получи ключ на https://aistudio.google.com"
    )
    if api_key_input:
        st.success("✓ Ключ введён")
    else:
        st.warning("Введи API ключ для работы")

    st.divider()
    st.markdown("**Что генерирует ИИ:**")
    for item in ["Полный текст заявки (7 разделов)", "AI-аудит: оценка 0–100", "3 стратегических совета", "DNSH & State Aid compliance", "Экспорт PDF + DOCX"]:
        st.markdown(f"&nbsp;&nbsp;✓ {item}")

    st.divider()
    st.caption("Язык определяется автоматически по выбранному гранту (UA / PL / EN)")


# ─── HEADER ───────────────────────────────────────────────────
st.markdown("# GrantMaster **Pro**")
st.markdown(
    '<p style="color:#607090; font-size:15px; margin-top:-12px;">Профессиональный AI-генератор грантовых заявок — EU / PL / UA</p>',
    unsafe_allow_html=True
)
st.markdown("---")

# ─── ШАГ 1: ВЫБОР ГРАНТА ──────────────────────────────────────
st.markdown("### Шаг 1 — Выберите грант")

cols = st.columns(3)
grant_names = list(GRANTS.keys())
for idx, name in enumerate(grant_names):
    g = GRANTS[name]
    col = cols[idx % 3]
    is_active = st.session_state.selected_grant == name
    card_class = "grant-card active" if is_active else "grant-card"
    with col:
        st.markdown(f"""
        <div class="{card_class}">
            <div class="grant-tag">{g['tag']}</div>
            <div class="grant-name">{name}</div>
            <div class="grant-amount">{g['amount']}</div>
            <div class="grant-desc">{g['desc']}</div>
        </div>
        """, unsafe_allow_html=True)
        btn_label = "✓ Выбран" if is_active else "Выбрать"
        if st.button(btn_label, key=f"select_{idx}", use_container_width=True,
                     type="primary" if is_active else "secondary"):
            st.session_state.selected_grant = name
            st.session_state.result_text = None
            st.session_state.prefill = {}
            st.rerun()

# ─── ШАГ 2: ФОРМА ─────────────────────────────────────────────
if st.session_state.selected_grant:
    grant_key = st.session_state.selected_grant
    grant_data = GRANTS[grant_key]
    pf = st.session_state.prefill

    st.markdown("---")
    st.markdown(f"### Шаг 2 — Данные проекта · `{grant_key}`")

    # Пример одной кнопкой вне формы
    col_ex, col_spacer = st.columns([2, 5])
    with col_ex:
        if st.button("📋 Вставить пример", help="Заполнит поля примером для этого гранта"):
            st.session_state.prefill = grant_data["example"].copy()
            st.rerun()

    st.markdown(" ")

    with st.form("main_form", clear_on_submit=False):
        col_left, col_right = st.columns([3, 2])

        with col_left:
            org = st.text_input(
                "Организация / Organization *",
                value=pf.get("org", ""),
                placeholder="Название + страна / город"
            )
            desc = st.text_area(
                "Описание проекта (идея + инновация) *",
                value=pf.get("desc", ""),
                height=150,
                placeholder="Что делаете, в чём инновация, какой impact..."
            )
            team = st.text_input(
                "Команда (роли + экспертиза)",
                value=pf.get("team", ""),
                placeholder="CEO + разработчики + эксперты"
            )

        with col_right:
            budget_val = int(pf.get("budget", 100000))
            budget = st.number_input(
                "Бюджет (EUR / PLN / UAH)",
                min_value=1000,
                max_value=50000000,
                value=budget_val,
                step=10000
            )
            st.markdown(" ")
            st.markdown(f"""
            <div style="background:#111826; border:1px solid #1E2A40; border-radius:10px; padding:16px;">
                <div style="font-size:11px; color:#607090; letter-spacing:0.06em; text-transform:uppercase; margin-bottom:8px;">Параметры генерации</div>
                <div style="font-size:12px; color:#8B9CC8; line-height:2;">
                    🌐 Язык заявки: <b style="color:#C9A84C">{grant_data['lang']}</b><br>
                    📋 Разделы: Summary, Excellence, Impact, Implementation, Budget, DNSH<br>
                    🤖 Модель: Gemini 1.5 Pro<br>
                    🌡 Temperature: 0.35 (юридический стиль)
                </div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown(" ")
        submitted = st.form_submit_button(
            "⚡ СГЕНЕРИРОВАТЬ ПРОФЕССИОНАЛЬНУЮ ЗАЯВКУ",
            use_container_width=True,
            type="primary"
        )

    # ─── ГЕНЕРАЦИЯ ────────────────────────────────────────────
    if submitted:
        if not api_key_input:
            st.error("Введи Gemini API ключ в боковой панели!")
        elif not org or not desc:
            st.error("Заполни обязательные поля: Организация и Описание проекта")
        else:
            try:
                genai.configure(api_key=api_key_input)

                with st.spinner("Gemini анализирует проект и готовит заявку..."):
                    progress = st.progress(0, text="Инициализация модели...")
                    import time

                    progress.progress(15, text="Анализ данных проекта...")
                    time.sleep(0.3)
                    progress.progress(30, text="Генерация AI-аудита...")

                    model = genai.GenerativeModel(
                        "gemini-1.5-pro-latest",
                        generation_config=genai.GenerationConfig(
                            temperature=0.35,
                            top_k=40,
                            top_p=0.9,
                            max_output_tokens=4096
                        )
                    )

                    prompt = f"""You are a Senior Grant Consultant with 15+ years of experience winning EU, Polish, and Ukrainian grants.

INPUT DATA:
- Grant Program: {grant_key}
- Organization: {org}
- Project Description: {desc}
- Team: {team if team else 'Not specified'}
- Budget: {budget:,} (EUR/PLN/UAH)

LANGUAGE RULES:
- If grant is UKRAINIAN (єРобота, МГО, UA programs) → write [APPLICATION] in UKRAINIAN
- If grant is POLISH (SMART FENG, PARP, NCBR) → write [APPLICATION] in POLISH
- Otherwise → write [APPLICATION] in ENGLISH
- Always write [SCORE] and [ADVICE] in ENGLISH

OUTPUT STRUCTURE (follow exactly, no markdown symbols, no asterisks):

[SCORE]
Score: X/100
Strengths: (2 concise sentences about what is strong)
Risks: (2 concise sentences about weak points)

[ADVICE]
1. (Strategic tip to improve funding chances — 2 sentences)
2. (Strategic tip to improve funding chances — 2 sentences)
3. (Strategic tip to improve funding chances — 2 sentences)

[APPLICATION]
1. EXECUTIVE SUMMARY
(3 paragraphs — clear value proposition, target beneficiaries, expected results)

2. PROJECT EXCELLENCE & INNOVATION
(2 paragraphs — TRL level, novelty, competitive advantage, state-of-the-art comparison)

3. IMPACT & SUSTAINABILITY
(2 paragraphs — quantified KPIs, social/economic impact, post-grant sustainability plan)

4. IMPLEMENTATION PLAN
(2 paragraphs — work packages, milestones, timeline, risk mitigation)

5. BUDGET JUSTIFICATION
(1 paragraph — personnel, equipment, subcontracting, cost-effectiveness rationale)

6. TEAM & CAPACITY
(1 paragraph — relevant expertise, track record, complementary skills)

7. DNSH & COMPLIANCE STATEMENT
(1 paragraph — Do No Significant Harm compliance, State Aid rules, GDPR if relevant)

Style: Formal legal/grant writing. No bullet points. Plain paragraphs. No markdown. Professional tone of official EU documentation."""

                    progress.progress(60, text="Генерация полного текста заявки...")
                    response = model.generate_content(prompt)
                    result_raw = response.text

                    progress.progress(90, text="Обработка результата...")
                    time.sleep(0.2)
                    progress.progress(100, text="Готово!")
                    time.sleep(0.3)
                    progress.empty()

                # Парсинг секций
                score_text = ""
                advice_text = ""
                app_text = result_raw

                if "[SCORE]" in result_raw and "[ADVICE]" in result_raw and "[APPLICATION]" in result_raw:
                    parts_a = result_raw.split("[ADVICE]", 1)
                    score_text = parts_a[0].replace("[SCORE]", "").strip()
                    parts_b = parts_a[1].split("[APPLICATION]", 1)
                    advice_text = parts_b[0].strip()
                    app_text = parts_b[1].strip()
                elif "[APPLICATION]" in result_raw:
                    parts = result_raw.split("[APPLICATION]", 1)
                    score_text = parts[0].replace("[SCORE]", "").replace("[ADVICE]", "").strip()
                    app_text = parts[1].strip()

                st.session_state.result_text = app_text
                st.session_state.score_text = score_text
                st.session_state.advice_text = advice_text
                st.session_state.org_name = org

            except Exception as e:
                st.error(f"Ошибка API: {str(e)}")
                st.info("Проверь API ключ и убедись, что Gemini 1.5 Pro доступен для твоего аккаунта.")

# ─── РЕЗУЛЬТАТ ────────────────────────────────────────────────
if st.session_state.result_text:
    st.markdown("---")
    st.markdown("### Результат")

    tab1, tab2, tab3 = st.tabs(["📄 Текст заявки", "📊 AI-аудит", "📥 Скачать"])

    # ── ТАБ 1: Заявка ──
    with tab1:
        st.markdown(f"""
        <div class="result-box">{st.session_state.result_text}</div>
        """, unsafe_allow_html=True)

    # ── ТАБ 2: Аудит ──
    with tab2:
        score_raw = st.session_state.score_text or ""
        advice_raw = st.session_state.advice_text or ""

        # Извлечь число из "Score: 78/100"
        import re
        score_match = re.search(r"Score:\s*(\d+)", score_raw)
        score_num = int(score_match.group(1)) if score_match else 72

        col_score, col_details = st.columns([1, 2])
        with col_score:
            color = "#4CAF50" if score_num >= 75 else "#C9A84C" if score_num >= 55 else "#E53935"
            st.markdown(f"""
            <div class="score-box">
                <div class="score-number" style="color:{color}">{score_num}</div>
                <div class="score-label">AI Score / 100</div>
                <div style="margin-top:12px; font-size:12px; color:#607090;">
                    {'🟢 Высокий потенциал' if score_num >= 75 else '🟡 Средний потенциал' if score_num >= 55 else '🔴 Требует доработки'}
                </div>
            </div>
            """, unsafe_allow_html=True)

        with col_details:
            if score_raw:
                cleaned = score_raw.replace("Score:", "").replace(f"{score_num}/100", "").strip()
                if cleaned:
                    st.markdown(f"""<div class="advice-item" style="border-left-color:#8B9CC8">{cleaned}</div>""",
                                unsafe_allow_html=True)

        if advice_raw:
            st.markdown(" ")
            st.markdown("**Стратегические рекомендации**")
            for i, line in enumerate(advice_raw.strip().split("\n")):
                line = line.strip()
                if line and len(line) > 10:
                    clean_line = re.sub(r"^\d+\.\s*", "", line)
                    if clean_line:
                        st.markdown(f'<div class="advice-item">{clean_line}</div>',
                                    unsafe_allow_html=True)

    # ── ТАБ 3: Скачать ──
    with tab3:
        org_safe = re.sub(r'[^a-zA-Z0-9_]', '_', st.session_state.org_name)[:40]
        grant_safe = re.sub(r'[^a-zA-Z0-9_]', '_', st.session_state.selected_grant or "Grant")[:30]
        timestamp = datetime.now().strftime("%Y%m%d")
        base_name = f"{org_safe}_{grant_safe}_{timestamp}"
        full_text = st.session_state.result_text

        # ── PDF ──
        try:
            class GrantPDF(FPDF):
                def header(self):
                    self.set_fill_color(11, 15, 26)
                    self.rect(0, 0, 210, 297, 'F')
                    self.set_draw_color(201, 168, 76)
                    self.set_line_width(0.8)
                    self.rect(10, 10, 190, 277)
                    self.ln(5)

                def footer(self):
                    self.set_y(-20)
                    self.set_font('Helvetica', 'I', 8)
                    self.set_text_color(96, 112, 144)
                    self.cell(0, 10, f'GrantMaster Pro · {st.session_state.selected_grant} · Page {self.page_no()}', align='C')

            pdf = GrantPDF()
            pdf.set_auto_page_break(auto=True, margin=25)
            pdf.add_page()
            pdf.set_margins(22, 22, 22)

            # Title block
            pdf.set_fill_color(17, 24, 38)
            pdf.rect(14, 22, 182, 28, 'F')
            pdf.set_font('Helvetica', 'B', 16)
            pdf.set_text_color(201, 168, 76)
            pdf.set_xy(14, 26)
            pdf.cell(182, 10, 'GRANT APPLICATION', align='C')
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(139, 156, 200)
            pdf.set_xy(14, 36)
            pdf.cell(182, 8, f'{st.session_state.selected_grant}  ·  {st.session_state.org_name}  ·  {datetime.now().strftime("%d.%m.%Y")}', align='C')
            pdf.ln(18)

            # Content
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(197, 205, 224)

            # Clean text
            clean = re.sub(r'[^\x00-\x7F\u00C0-\u024F\u0100-\u017E]', '', full_text)
            paragraphs = clean.split('\n')

            for para in paragraphs:
                para = para.strip()
                if not para:
                    pdf.ln(3)
                    continue
                # Section header detection (all caps or starts with digit+dot)
                is_header = (para.upper() == para and len(para) > 3) or re.match(r'^\d+\.\s+[A-Z]', para)
                if is_header:
                    pdf.ln(4)
                    pdf.set_font('Helvetica', 'B', 11)
                    pdf.set_text_color(201, 168, 76)
                    pdf.multi_cell(0, 6, para)
                    pdf.set_font('Helvetica', '', 10)
                    pdf.set_text_color(197, 205, 224)
                    pdf.ln(2)
                else:
                    pdf.multi_cell(0, 5.5, para)
                    pdf.ln(1)

            pdf_bytes = bytes(pdf.output())

            # ── DOCX ──
            doc = Document()
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Document settings
            section = doc.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
            section.top_margin = Inches(1.1)
            section.bottom_margin = Inches(1.1)

            # Title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_t = title_para.add_run("GRANT APPLICATION")
            run_t.bold = True
            run_t.font.size = Pt(18)
            run_t.font.color.rgb = RGBColor(0x00, 0x30, 0x87)

            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_s = sub_para.add_run(f"{st.session_state.selected_grant}  ·  {st.session_state.org_name}  ·  {datetime.now().strftime('%d.%m.%Y')}")
            run_s.font.size = Pt(10)
            run_s.font.color.rgb = RGBColor(0x60, 0x70, 0x90)

            doc.add_paragraph()

            for para in full_text.split('\n'):
                para = para.strip()
                if not para:
                    doc.add_paragraph()
                    continue
                is_header = (para.upper() == para and len(para) > 3) or re.match(r'^\d+\.\s+[A-Z]', para)
                p = doc.add_paragraph()
                run = p.add_run(para)
                if is_header:
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)
                else:
                    run.font.size = Pt(10.5)

            docx_bio = io.BytesIO()
            doc.save(docx_bio)
            docx_bio.seek(0)

            # Download buttons
            st.markdown(" ")
            dcol1, dcol2 = st.columns(2)
            with dcol1:
                st.download_button(
                    "📕  Скачать PDF",
                    data=pdf_bytes,
                    file_name=f"{base_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            with dcol2:
                st.download_button(
                    "📄  Скачать DOCX (Word)",
                    data=docx_bio.getvalue(),
                    file_name=f"{base_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            st.markdown(" ")
            st.caption(f"Файл: `{base_name}.pdf / .docx`  ·  Сгенерировано {datetime.now().strftime('%d.%m.%Y %H:%M')}")

        except Exception as e:
            st.error(f"Ошибка при создании файлов: {str(e)}")
            # Fallback: plain text download
            st.download_button(
                "📥 Скачать TXT (fallback)",
                data=full_text.encode('utf-8'),
                file_name=f"{base_name}.txt",
                mime="text/plain"
            )

# ─── FOOTER ───────────────────────────────────────────────────
st.markdown("---")
st.markdown(
    '<p style="text-align:center; color:#2A3A58; font-size:12px;">© GrantMaster Pro 2026 · AI-powered · EU / PL / UA · Удали API ключ после публичного деплоя</p>',
    unsafe_allow_html=True
)
